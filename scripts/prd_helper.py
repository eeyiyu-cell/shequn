# -*- coding: utf-8 -*-
"""在 skill prd_docx.py 基础上扩展：把截图嵌入「需求详情」三列表的原型图列。"""
import os, sys
from PIL import Image
from docx.shared import Cm, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT

SKILL = r"C:\Users\Admin\.claude\skills\phprd2\scripts"
sys.path.insert(0, SKILL)
from prd_docx import (PrdDoc, _set_run_font, _shade, _table_borders,
                      CHAMPAGNE_HEADER, HEADER_TEXT, BODY_TEXT, BODY_PT)

DOC_IMG_DIR = None  # 由 build 脚本设置

class PrdBuilder(PrdDoc):
    def page_table(self, rows, img_dir, img_w_cm=5.6, widths_cm=(6.0, 3.2, 8.6)):
        """需求详情三列表：原型图(内嵌截图) | 功能 | 逻辑说明(▪ 分行)。
        rows: [(img_basename_no_ext, 功能名, 逻辑说明多行文本), ...]
        img 为 img_dir 下的 <basename>.jpg；缺图则写占位文字。
        """
        headers = ["原型图", "功能", "逻辑说明"]
        t = self.doc.add_table(rows=1, cols=3)
        t.alignment = WD_TABLE_ALIGNMENT.LEFT
        _table_borders(t)
        for i, h in enumerate(headers):
            cell = t.rows[0].cells[i]
            _shade(cell, CHAMPAGNE_HEADER)
            run = cell.paragraphs[0].add_run(h)
            _set_run_font(run, BODY_PT, bold=True, color=HEADER_TEXT)
        for basename, func, logic in rows:
            cells = t.add_row().cells
            # 原型图列：内嵌图
            imgp = os.path.join(img_dir, basename + ".jpg")
            p0 = cells[0].paragraphs[0]
            if os.path.exists(imgp):
                run = p0.add_run()
                run.add_picture(imgp, width=Cm(img_w_cm))
            else:
                r = p0.add_run("〔待补：%s 原型图〕" % func)
                _set_run_font(r, BODY_PT, color=BODY_TEXT)
            # 功能列
            p1 = cells[1].paragraphs[0]
            r1 = p1.add_run(func)
            _set_run_font(r1, BODY_PT, bold=True, color=BODY_TEXT)
            # 逻辑说明列：\n 分行
            lines = str(logic).split("\n")
            p2 = cells[2].paragraphs[0]
            rr = p2.add_run(lines[0])
            _set_run_font(rr, BODY_PT, color=BODY_TEXT)
            for extra in lines[1:]:
                pe = cells[2].add_paragraph()
                re = pe.add_run(extra)
                _set_run_font(re, BODY_PT, color=BODY_TEXT)
        for i, w in enumerate(widths_cm):
            for row in t.rows:
                row.cells[i].width = Cm(w)
        return t

if __name__ == "__main__":
    # 自测：单行三列表内嵌一张既有图，验证图入单元格机制
    b = PrdBuilder()
    b.title("【自测】图入单元格")
    b.h2("2.2.1 测试")
    v2doc = r"C:\Users\Admin\Desktop\身心灵\原型\红魔方社群积分系统\output\prd_assets_v2\_doc"
    b.page_table([("adm-01-活动管理列表", "活动管理列表",
                   "▪ 路径：Power Club+ → 活动管理\n▪ 测试第二行")], v2doc)
    out = r"C:\Users\Admin\AppData\Local\Temp\claude\C--Users-Admin\5c8f5f13-1098-4076-9d24-28d84c193046\scratchpad\selftest.docx"
    b.save(out)
    # 读回验证
    from docx import Document
    d = Document(out)
    imgs = len(d.inline_shapes)
    print("saved", out, "inline_shapes=", imgs, "tables=", len(d.tables))
